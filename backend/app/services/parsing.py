"""Resume ingestion: bytes in, structured ParsedResume out."""

from __future__ import annotations

import io
import logging

from ..schemas import ParsedResume, ResumeBullet
from . import textkit
from .skills import extract_skills

log = logging.getLogger("fitscope.parsing")

MIN_USEFUL_CHARS = 200


def _pdf_to_text(data: bytes) -> tuple[str, int, list[str]]:
    import pdfplumber

    warnings: list[str] = []
    pages_text: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=1.5, y_tolerance=3) or ""
            if not text.strip():
                # Two-column or graphic-heavy layouts often need word grouping.
                words = page.extract_words() or []
                text = " ".join(w["text"] for w in words)
            pages_text.append(text)
    joined = "\n".join(pages_text)
    if len(joined.strip()) < MIN_USEFUL_CHARS:
        warnings.append(
            "Very little text could be extracted. The PDF is likely a scan or image export, "
            "which most ATS parsers also fail to read."
        )
    return joined, page_count, warnings


def _docx_to_text(data: bytes) -> tuple[str, list[str]]:
    import docx  # python-docx

    warnings: list[str] = []
    document = docx.Document(io.BytesIO(data))
    lines: list[str] = []
    for para in document.paragraphs:
        style = (para.style.name or "").lower()
        text = para.text.strip()
        if not text:
            continue
        lines.append(f"- {text}" if "list" in style and not text.startswith("-") else text)
    if document.tables:
        warnings.append(
            "Resume uses tables for layout. Several ATS parsers read table cells out of order."
        )
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(dict.fromkeys(cells)))
    return "\n".join(lines), warnings


def parse_resume_bytes(data: bytes, filename: str) -> ParsedResume:
    name = (filename or "").lower()
    warnings: list[str] = []
    pages: int | None = None

    if name.endswith(".pdf") or data[:5] == b"%PDF-":
        kind = "pdf"
        raw, pages, warnings = _pdf_to_text(data)
    elif name.endswith((".docx", ".doc")) or data[:2] == b"PK":
        kind = "docx"
        raw, warnings = _docx_to_text(data)
        if name.endswith(".doc") and not name.endswith(".docx"):
            warnings.append("Legacy .doc detected; export as PDF or .docx for reliable parsing.")
    else:
        kind = "text"
        raw = data.decode("utf-8", errors="replace")

    parsed = parse_resume_text(raw, source_kind=kind, pages=pages)
    parsed.parse_warnings = warnings + parsed.parse_warnings
    return parsed


def parse_resume_text(
    raw: str, source_kind: str = "text", pages: int | None = None
) -> ParsedResume:
    text = textkit.clean_text(raw)
    sections = textkit.split_sections(text)
    bullets = textkit.extract_bullets(sections)
    emails, phones, links = textkit.extract_contacts(text)

    warnings: list[str] = []
    if not bullets:
        warnings.append("No achievement lines were detected. Check bullet formatting.")

    return ParsedResume(
        raw_text=text,
        char_count=len(text),
        word_count=len(text.split()),
        bullets=[ResumeBullet(**b) for b in bullets],
        sections=sections,
        emails=emails,
        phones=phones,
        links=links,
        skills=extract_skills(text),
        years_experience=textkit.estimate_years(text),
        source_kind=source_kind,  # type: ignore[arg-type]
        pages=pages,
        parse_warnings=warnings,
    )
